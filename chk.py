import json, sys, os, re
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

fn = '计算机专业英语第一次作业（硬件知识）答案+翻译+解析.docx'
doc = Document(fn)

# 收集所有非空段落
paras = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        paras.append((i, t))

print(f'原始 {len(doc.paragraphs)} 段 → 非空 {len(paras)} 段\n')

# 找结构段
for idx, (i, t) in enumerate(paras):
    # 大题标题
    if re.match(r'^[一二三四五六七八九十]、', t) or '答题' in t or '单选题' in t or '多选题' in t or '判断题' in t or '简答题' in t or '翻译题' in t or '阅读理解' in t:
        print(f'\n[{idx}] p{i}: ### {t[:120]}')
        continue
    # 题号
    m = re.match(r'^(\d{1,3})[.)、]\s*(.+)', t)
    if m:
        print(f'[{idx}] p{i}: Q{m.group(1)}: {t[:120]}')
    elif '答案' in t[:10]:
        print(f'[{idx}] p{i}: ANS: {t[:120]}')
    elif '解析' in t[:10] or '分析' in t[:10]:
        print(f'[{idx}] p{i}: EXP: {t[:120]}')
    elif '翻译' in t[:10] or '中文' in t[:10]:
        print(f'[{idx}] p{i}: TRANS: {t[:120]}')
    elif len(t) > 10 and not t.startswith('A') and not t.startswith('B') and not t.startswith('C') and not t.startswith('D') and not t.startswith('E'):
        # 可能是选项或其他
        if re.match(r'^[A-E][.)、\s]', t):
            print(f'[{idx}] p{i}: OPT: {t[:120]}')
        else:
            print(f'[{idx}] p{i}: TXT: {t[:120]}')
    else:
        print(f'[{idx}] p{i}: OPT: {t[:120]}')
